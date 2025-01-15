import fitz  # PyMuPDF
import json
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime
import uuid
import os
from docx import Document  # Add this import at the top

class DocumentFieldAutomation:
    def __init__(self):
        self.field_mappings = {}
        self.app_folder = None
        self.templates_folder = None
        self.cases_folder = None
        self.config_file = None

    def initialize_app_structure(self):
        """Set up or load the application folder structure"""
        self.app_folder = Path("/Users/206753750/Desktop/Lauris")
        
        print(f"Using application folder: {self.app_folder}")
        
        if not self.app_folder.exists():
            messagebox.showerror("Error", f"Folder not found: {self.app_folder}")
            return False
            
        self.templates_folder = self.app_folder / "templates"
        self.cases_folder = self.app_folder / "cases"
        self.config_file = self.app_folder / "field_mappings.json"
        
        self.templates_folder.mkdir(exist_ok=True)
        self.cases_folder.mkdir(exist_ok=True)
        
        pdf_templates = list(self.templates_folder.glob('*.pdf'))
        
        if not self.config_file.exists() or not pdf_templates:
            messagebox.showinfo(
                "Setup Required",
                f"Please place template documents in:\n{self.templates_folder}\n"
                f"And configuration file in:\n{self.config_file}"
            )
            return False
            
        return True

    def load_configuration(self):
        """Load field mapping configuration from JSON file"""
        try:
            with open(self.config_file, 'r', encoding='utf-8') as f:
                self.field_mappings = json.load(f)
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load configuration: {str(e)}")
            return False

    def add_text_to_pdf(self, input_path, output_path, text, page_number, x, y):
        """Add text to PDF using PyMuPDF"""
        try:
            print(f"\nAdding text to PDF:")
            print(f"File: {input_path}")
            print(f"Text: {text}")
            print(f"Page: {page_number}, Position: ({x}, {y})")
            
            # Create a temporary file path
            temp_path = str(output_path) + ".temp"
            
            # First copy the original file to temp location
            with open(input_path, 'rb') as src, open(temp_path, 'wb') as dst:
                dst.write(src.read())
            
            # Open the temp file
            doc = fitz.open(temp_path)
            page = doc[page_number]
            
            # Get page dimensions
            rect = page.rect
            print(f"Page dimensions: {rect}")
            print(f"Page rotation: {page.rotation}")
            
            # For landscape pages (where width > height), adjust coordinates
            is_landscape = rect.width > rect.height
            if is_landscape:
                print("Landscape orientation detected")
                # Adjust coordinates for landscape
                final_x = x
                final_y = y
            else:
                print("Portrait orientation detected")
                final_x = x
                final_y = y
                
            # Insert text
            text_point = fitz.Point(final_x, final_y)
            page.insert_text(text_point, text, fontname="helv", fontsize=11, color=(0, 0, 0))
            
            # Save to the final output location
            doc.save(str(output_path))
            doc.close()
            
            # Remove temp file
            os.remove(temp_path)
            
            print("Successfully added text to PDF")
            return True
            
        except Exception as e:
            print(f"Error adding text: {str(e)}")
            return False

    def create_case_folder(self, case_name):
        """Create new folder for the case with unique identifier"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        case_folder = self.cases_folder / f"{timestamp}_{unique_id}_{case_name}"
        case_folder.mkdir(parents=True, exist_ok=True)
        return case_folder

    def fill_word(self, input_path, output_path, text_replacements=None, table_locations=None):
        """Fill Word document fields with text"""
        try:
            print(f"\nProcessing Word document:")
            print(f"Input: {input_path}")
            print(f"Output: {output_path}")
            print(f"Table locations: {table_locations}")
            
            # Load the document
            doc = Document(input_path)
            
            # Track if we made any changes
            changes_made = False
            
            # Process tables according to specified locations
            if table_locations:
                print("\nProcessing table locations:")
                for location in table_locations:
                    table_idx = location.get('table_index', 0)
                    row_idx = location.get('row', 0)
                    col_idx = location.get('column', 0)
                    
                    print(f"Looking for table {table_idx}, row {row_idx}, column {col_idx}")
                    
                    try:
                        # Get the specified table
                        if table_idx < len(doc.tables):
                            table = doc.tables[table_idx]
                            print(f"Found table {table_idx}")
                            
                            # Get the specified cell
                            if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                print(f"Found cell at row {row_idx}, column {col_idx}")
                                print(f"Current cell content: {cell.text}")
                                
                                # Update cell content
                                if not cell.text.strip():  # Check if cell is empty
                                    text = list(text_replacements.values())[0]  # Get the replacement text
                                    cell.text = text
                                    changes_made = True
                                    print(f"Updated cell with text: {text}")
                                else:
                                    print("Cell is not empty, skipping")
                            else:
                                print(f"Row {row_idx} or column {col_idx} index out of range")
                        else:
                            print(f"Table index {table_idx} out of range")
                            
                    except Exception as e:
                        print(f"Error processing table location: {str(e)}")
            
            # Save the document
            doc.save(output_path)
            print("Successfully saved Word document")
            return changes_made
            
        except Exception as e:
            print(f"Error processing Word document: {str(e)}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")
            return False

    def process_documents(self, case_folder):
        """Process all documents with user input"""
        user_input = {
            field_name: entry.get() 
            for field_name, entry in self.inputs.items()
        }
        
        print("\nDEBUG - Starting document processing")
        print(f"Case folder: {case_folder}")
        print(f"User input: {user_input}")
        
        # Process each document
        for template in self.templates_folder.iterdir():
            file_extension = template.suffix.lower()
            output_path = case_folder / template.name
            
            # Copy the original file first
            with open(template, 'rb') as src, open(output_path, 'wb') as dst:
                dst.write(src.read())
            
            if file_extension == '.pdf':
                print(f"\nProcessing PDF: {template.name}")
                # Process PDF files as before
                modified = False
                for field in self.field_mappings['fields']:
                    if template.name in field['locations']:
                        text = user_input[field['name']]
                        print(f"Adding field: {field['name']} = {text}")
                        
                        for page_num, x, y in field['locations'][template.name]:
                            success = self.add_text_to_pdf(
                                str(output_path),
                                str(output_path),
                                text,
                                page_num,
                                x,
                                y
                            )
                            if success:
                                modified = True
                                print(f"Successfully added text to {template.name}")
                            else:
                                print(f"Failed to add text to {template.name}")
                
                if not modified:
                    print(f"No modifications made to {template.name}")
                    
            elif file_extension == '.docx':
                print(f"\nProcessing Word document: {template.name}")
                # Process Word documents
                replacements = {}
                for field in self.field_mappings['fields']:
                    if template.name in field['locations']:
                        text = user_input[field['name']]
                        # Get placeholder texts for this document
                        for placeholder in field['locations'][template.name]:
                            replacements[placeholder] = text
                
                if replacements:
                    success = self.fill_word(output_path, output_path, replacements)
                    if success:
                        print(f"Successfully modified {template.name}")
                    else:
                        print(f"Failed to modify {template.name}")
                else:
                    print(f"No modifications needed for {template.name}")

    def create_input_form(self):
        """Create GUI form for user input"""
        self.root = tk.Tk()
        self.root.title("Document Field Automation")
        
        case_frame = ttk.LabelFrame(self.root, text="Case Information", padding=10)
        case_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(case_frame, text="Case Name:").pack()
        self.case_name_entry = ttk.Entry(case_frame)
        self.case_name_entry.pack(pady=5)
        
        fields_frame = ttk.LabelFrame(self.root, text="Document Fields", padding=10)
        fields_frame.pack(fill="x", padx=10, pady=5)
        
        self.inputs = {}
        for field in self.field_mappings['fields']:
            label = ttk.Label(fields_frame, text=field['display_name'])
            label.pack(pady=5)
            
            entry = ttk.Entry(fields_frame)
            entry.pack(pady=5)
            self.inputs[field['name']] = entry
        
        self.status_var = tk.StringVar()
        status_label = ttk.Label(self.root, textvariable=self.status_var)
        status_label.pack(pady=5)
        
        submit_btn = ttk.Button(self.root, text="Process Documents", 
                              command=self.validate_and_process)
        submit_btn.pack(pady=20)
        
        info_frame = ttk.LabelFrame(self.root, text="System Information", padding=10)
        info_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(info_frame, text=f"Templates Folder: {self.templates_folder}").pack()
        ttk.Label(info_frame, text=f"Cases Folder: {self.cases_folder}").pack()
        
        self.root.mainloop()

    def validate_and_process(self):
        """Validate inputs and process documents"""
        case_name = self.case_name_entry.get().strip()
        if not case_name:
            self.status_var.set("Please enter a case name!")
            return
        
        try:
            case_folder = self.create_case_folder(case_name)
            self.process_documents(case_folder)
            self.status_var.set(
                f"Documents processed successfully!\n"
                f"Saved in: {case_folder.name}"
            )
            
            # Ask what to do next
            choice = messagebox.askyesno(
                "What Next?",
                "Would you like to process another case?\n\n'Yes' - Process another case\n'No' - Close the application",
                icon='question'
            )
            
            if choice:
                # Clear all inputs for new case
                self.case_name_entry.delete(0, tk.END)
                for entry in self.inputs.values():
                    entry.delete(0, tk.END)
                self.status_var.set("Ready for new case")
            else:
                # Close the application
                self.root.quit()
                self.root.destroy()
                
        except Exception as e:
            print(f"\nERROR in validate_and_process: {str(e)}")
            import traceback
            print(traceback.format_exc())
            self.status_var.set(f"Error processing documents: {str(e)}")

def main():
    app = DocumentFieldAutomation()
    
    if not app.initialize_app_structure():
        return
        
    if not app.load_configuration():
        return
        
    app.create_input_form()

if __name__ == "__main__":
    main()