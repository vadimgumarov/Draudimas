# src/word/word_reader.py
from pathlib import Path
from docx import Document
import shutil
from datetime import datetime

class WordHandler:
    @staticmethod
    def read_docx(docx_path: Path) -> bool:
        """Read a Word document and verify it can be opened."""
        try:
            if not docx_path.exists():
                print(f"Error: File not found: {docx_path}")
                return False
                
            doc = Document(docx_path)
            print(f"Successfully opened Word document: {docx_path.name}")
            print(f"Number of tables: {len(doc.tables)}")
            return True
                
        except Exception as e:
            print(f"Error: {str(e)}")
            return False

    @staticmethod
    def add_text_to_docx(template_path: Path, case_folder: Path, text: str, table: int, row: int, col: int) -> bool:
        """Add text to Word document at specified table cell."""
        try:
            output_path = case_folder / template_path.name
            
            # If this is the first time adding text to this case, copy the template
            if not output_path.exists():
                case_folder.mkdir(parents=True, exist_ok=True)
                shutil.copy2(str(template_path), str(output_path))
            
            # Open and modify the document
            doc = Document(output_path)
            
            # Verify table exists
            if table >= len(doc.tables):
                raise Exception(f"Table index {table} out of range (document has {len(doc.tables)} tables)")
            
            table_obj = doc.tables[table]
            
            # Verify row exists
            if row >= len(table_obj.rows):
                raise Exception(f"Row {row} out of range (table has {len(table_obj.rows)} rows)")
            
            # Verify column exists
            if col >= len(table_obj.rows[row].cells):
                raise Exception(f"Column {col} out of range (row has {len(table_obj.rows[row].cells)} cells)")
            
            # Add text to specified cell while preserving existing content
            cell = table_obj.rows[row].cells[col]
            existing_text = cell.text.strip()
            
            # Add a space between existing text and new text if there is existing text
            if existing_text:
                new_text = f"{existing_text} {text}"
            else:
                new_text = text
                
            cell.text = new_text
            
            # Save the document
            doc.save(str(output_path))
            
            return True
            
        except Exception as e:
            print(f"Error: {str(e)}")
            return False

    @staticmethod
    def add_text_to_docx_multi(template_path: Path, case_folder: Path, text: str, coordinates_list) -> bool:
        """Add text to Word document at multiple specified locations."""
        try:
            # If coordinates_list is a single dictionary, convert it to a list
            if isinstance(coordinates_list, dict):
                coordinates_list = [coordinates_list]

            # Process each coordinate set
            for coord_info in coordinates_list:
                success = WordHandler.add_text_to_docx(
                    template_path,
                    case_folder,
                    text,
                    coord_info["table"],
                    coord_info["row"],
                    coord_info["col"]
                )
                if not success:
                    print(f"Warning: Failed to add text at coordinates: {coord_info}")

            return True

        except Exception as e:
            print(f"Error in multi-coordinate text addition: {str(e)}")
            return False

    @staticmethod
    def create_case_folder(base_dir: Path, case_name: str) -> Path:
        """Create a case folder."""
        try:
            case_folder = base_dir / case_name
            print(f"Debug: Creating case folder: {case_folder}")
            case_folder.mkdir(parents=True, exist_ok=True)
            print(f"Debug: Case folder created/verified successfully")
            return case_folder
        except Exception as e:
            print(f"Debug: Error creating case folder: {str(e)}")
            raise