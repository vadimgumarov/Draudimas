from docx import Document

def analyze_word_doc(filename):
    print(f"\nAnalyzing: {filename}")
    doc = Document(filename)
    
    print(f"Number of tables: {len(doc.tables)}")
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx}:")
        print(f"Rows: {len(table.rows)}, Columns: {len(table.rows[0].cells)}")
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if cell.text.strip():  # Only print non-empty cells
                    print(f"Row {row_idx}, Column {col_idx}: {cell.text.strip()}")

# Analyze both Word documents
analyze_word_doc("Europos_paramos_KPP.docx")
analyze_word_doc("NAC_2025.docx")
